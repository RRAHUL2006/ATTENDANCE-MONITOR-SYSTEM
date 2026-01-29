import os, sys
from flask import Flask, render_template_string, request, redirect, flash
from werkzeug.security import generate_password_hash, check_password_hash
from flask import get_flashed_messages
import sqlite3
from datetime import date
import pandas as pd
from flask import session
from openpyxl import Workbook
from flask import send_file
import tempfile
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
import json
import webbrowser
import threading
import docx
from PyPDF2 import PdfReader
import io


def init_db_if_needed():
    conn = get_db()
    c = conn.cursor()

    c.execute("""
    CREATE TABLE IF NOT EXISTS admin (
        password TEXT
    )
    """)

    c.execute("""
    CREATE TABLE IF NOT EXISTS classes (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        class_name TEXT UNIQUE
    )
    """)

    c.execute("""
    CREATE TABLE IF NOT EXISTS students (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        roll_no TEXT,
        name TEXT,
        class_id INTEGER,
        FOREIGN KEY(class_id) REFERENCES classes(id)
    )
    """)

    c.execute("""
    CREATE TABLE IF NOT EXISTS attendance (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        student_id INTEGER,
        date TEXT,
        status TEXT,
        od_reason TEXT,
        FOREIGN KEY(student_id) REFERENCES students(id)
    )
    """)

   


    conn.commit()
    conn.close()


def resource_path(relative_path):
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)


def get_attendance_status(percent):
    """Return status badge and color for attendance percentage"""
    if percent >= 75:
        return {"status": "Safe", "badge": "🟢", "color": "#28a745", "class": "safe"}
    elif percent >= 65:
        return {"status": "Warning", "badge": "🟡", "color": "#ffc107", "class": "warning"}
    else:
        return {"status": "Defaulter", "badge": "🔴", "color": "#dc3545", "class": "defaulter"}


def get_student_attendance_stats(student_id, db):
    """Get detailed attendance stats for a student"""
    total = db.execute(
        "SELECT COUNT(*) FROM attendance WHERE student_id=?",
        (student_id,)
    ).fetchone()[0]

    present = db.execute(
        "SELECT COUNT(*) FROM attendance WHERE student_id=? AND status IN ('Present','OD')",
        (student_id,)
    ).fetchone()[0]

    absent = total - present
    percent = round((present / total) * 100, 2) if total else 0
    status_info = get_attendance_status(percent)

    return {
        "total": total,
        "present": present,
        "absent": absent,
        "percent": percent,
        "status": status_info["status"],
        "badge": status_info["badge"],
        "color": status_info["color"],
        "class": status_info["class"]
    }


ROLL_KEYS = [
    "roll", "rollno", "roll_no",
    "reg", "regno", "register", "registerno",
    "registration", "registrationno"
]

NAME_KEYS = [
    "name", "studentname", "student_name",
    "fullname", "full_name"
]


def normalize(col):
    return (
        col.lower()
           .replace(" ", "")
           .replace("_", "")
           .replace("-", "")
    )


def detect_columns(columns):
    roll_col = None
    name_col = None

    for col in columns:
        n = normalize(col)

        if not roll_col and any(k in n for k in ROLL_KEYS):
            roll_col = col

        if not name_col and any(k in n for k in NAME_KEYS):
            name_col = col

    return roll_col, name_col


app = Flask(__name__, static_folder=resource_path("static"))
app.secret_key = "attendance_secret"


def get_db():
    appdata = os.path.join(os.getenv("APPDATA"), "PerfectAttendance")
    os.makedirs(appdata, exist_ok=True)

    db_path = os.path.join(appdata, "database.db")

    if not os.path.exists(db_path):
        src = resource_path("database.db")
        if os.path.exists(src):
            import shutil
            shutil.copy(src, db_path)
        else:
            open(db_path, "w").close()

    return sqlite3.connect(db_path)


# ====================== PREMIUM STYLES & SCRIPTS ======================
BASE_STYLES = """
<style>
    * {
        margin: 0;
        padding: 0;
        box-sizing: border-box;
    }

    html, body {
        height: 100%;
    }

    body {
        font-family: 'Segoe UI', 'Roboto', 'Helvetica Neue', sans-serif;
        background: linear-gradient(135deg, #667eea 0%, #764ba2 50%, #f093fb 100%);
        min-height: 100vh;
        padding: 20px;
        color: #333;
    }

    .container {
        max-width: 1400px;
        margin: 0 auto;
    }

    .header {
        background: rgba(255, 255, 255, 0.98);
        backdrop-filter: blur(20px);
        border-radius: 25px;
        padding: 45px;
        margin-bottom: 35px;
        box-shadow: 0 25px 80px rgba(0, 0, 0, 0.15);
        display: flex;
        align-items: center;
        justify-content: space-between;
        flex-wrap: wrap;
        gap: 30px;
        border: 1px solid rgba(255, 255, 255, 0.5);
    }

    .header-left {
        display: flex;
        align-items: center;
        gap: 25px;
        flex: 1;
        min-width: 300px;
    }

    .header-logo {
        width: 140px;
        height: 140px;
        border-radius: 20px;
        overflow: hidden;
        box-shadow: 0 15px 50px rgba(102, 126, 234, 0.4);
        border: 4px solid #667eea;
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        display: flex;
        align-items: center;
        justify-content: center;
        font-size: 3em;
    }

    .header-logo img {
        width: 100%;
        height: 100%;
        object-fit: cover;
    }

    .header-title {
        flex: 1;
    }

    .header-title h1 {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        background-clip: text;
        font-size: 2.8em;
        margin-bottom: 8px;
        font-weight: 800;
        letter-spacing: -1px;
    }

    .header-title p {
        color: #666;
        font-size: 1.15em;
        font-weight: 500;
    }

    .header-actions {
        display: flex;
        gap: 12px;
        align-items: center;
    }

    .btn {
        padding: 14px 30px;
        border: none;
        border-radius: 12px;
        cursor: pointer;
        font-size: 1em;
        font-weight: 700;
        transition: all 0.4s cubic-bezier(0.175, 0.885, 0.32, 1.275);
        text-decoration: none;
        display: inline-flex;
        align-items: center;
        gap: 10px;
        box-shadow: 0 10px 30px rgba(0, 0, 0, 0.1);
        border: 2px solid transparent;
    }

    .btn-primary {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
    }

    .btn-primary:hover {
        transform: translateY(-4px);
        box-shadow: 0 20px 50px rgba(102, 126, 234, 0.5);
    }

    .btn-secondary {
        background: #f8f9fa;
        color: #667eea;
        border: 2px solid #667eea;
    }

    .btn-secondary:hover {
        background: #667eea;
        color: white;
        transform: translateY(-4px);
    }

    .btn-danger {
        background: linear-gradient(135deg, #dc3545 0%, #c82333 100%);
        color: white;
    }

    .btn-danger:hover {
        transform: translateY(-4px);
        box-shadow: 0 20px 50px rgba(220, 53, 69, 0.5);
    }

    .add-class-form {
        background: rgba(255, 255, 255, 0.98);
        border-radius: 20px;
        padding: 35px;
        margin-bottom: 35px;
        box-shadow: 0 20px 60px rgba(0, 0, 0, 0.1);
        border: 1px solid rgba(255, 255, 255, 0.5);
    }

    .add-class-form h2 {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        background-clip: text;
        margin-bottom: 25px;
        font-size: 1.9em;
        font-weight: 800;
    }

    .form-group {
        display: flex;
        gap: 12px;
        align-items: center;
    }

    .form-group input {
        flex: 1;
        padding: 14px 18px;
        border: 2px solid #e8e8e8;
        border-radius: 12px;
        font-size: 1em;
        transition: all 0.3s ease;
        font-weight: 500;
    }

    .form-group input:focus {
        outline: none;
        border-color: #667eea;
        box-shadow: 0 0 20px rgba(102, 126, 234, 0.3);
        background: #f8faff;
    }

    .classes-grid {
        display: grid;
        grid-template-columns: repeat(auto-fill, minmax(360px, 1fr));
        gap: 30px;
        margin-bottom: 40px;
    }

    .class-card {
        background: rgba(255, 255, 255, 0.98);
        backdrop-filter: blur(10px);
        border-radius: 18px;
        padding: 35px;
        box-shadow: 0 20px 60px rgba(0, 0, 0, 0.1);
        transition: all 0.4s cubic-bezier(0.175, 0.885, 0.32, 1.275);
        border: 2px solid transparent;
        position: relative;
        overflow: hidden;
        animation: slideUp 0.6s ease forwards;
    }

    .class-card:nth-child(1) { animation-delay: 0.1s; }
    .class-card:nth-child(2) { animation-delay: 0.2s; }
    .class-card:nth-child(3) { animation-delay: 0.3s; }
    .class-card:nth-child(4) { animation-delay: 0.4s; }

    @keyframes slideUp {
        from {
            opacity: 0;
            transform: translateY(30px);
        }
        to {
            opacity: 1;
            transform: translateY(0);
        }
    }

    .class-card:hover {
        transform: translateY(-12px);
        box-shadow: 0 35px 80px rgba(102, 126, 234, 0.3);
        border-color: #667eea;
    }

    .class-card::before {
        content: '';
        position: absolute;
        top: 0;
        left: 0;
        right: 0;
        height: 5px;
        background: linear-gradient(90deg, #667eea 0%, #764ba2 50%, #f093fb 100%);
    }

    .class-header {
        display: flex;
        justify-content: space-between;
        align-items: center;
        margin-bottom: 25px;
        position: relative;
        z-index: 1;
    }

    .class-name {
        font-size: 1.6em;
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        background-clip: text;
        font-weight: 800;
    }

    .class-badge {
        display: inline-block;
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        padding: 8px 16px;
        border-radius: 20px;
        font-size: 0.9em;
        font-weight: 700;
        box-shadow: 0 5px 20px rgba(102, 126, 234, 0.3);
    }

    .class-stats {
        background: linear-gradient(135deg, #f8faff 0%, #f0f4ff 100%);
        border-radius: 15px;
        padding: 25px;
        margin-bottom: 25px;
        position: relative;
        z-index: 1;
        border-left: 4px solid #667eea;
    }

    .stat-item {
        display: flex;
        justify-content: space-between;
        margin-bottom: 15px;
        font-size: 1.05em;
        color: #555;
        font-weight: 500;
    }

    .stat-item:last-child {
        margin-bottom: 0;
    }

    .stat-value {
        font-weight: 800;
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        background-clip: text;
    }

    .stat-defaulter {
        color: #dc3545;
        font-weight: 800;
    }

    .class-actions {
        display: flex;
        gap: 10px;
        margin-top: 25px;
        position: relative;
        z-index: 1;
    }

    .class-actions .btn {
        flex: 1;
        padding: 12px 15px;
        font-size: 0.95em;
        justify-content: center;
    }

    .status-badge {
        display: inline-block;
        padding: 10px 16px;
        border-radius: 20px;
        font-size: 0.95em;
        font-weight: 700;
        margin-left: 10px;
        box-shadow: 0 5px 15px rgba(0, 0, 0, 0.1);
    }

    .status-safe {
        background: linear-gradient(135deg, #d4edda 0%, #c3e6cb 100%);
        color: #155724;
    }

    .status-warning {
        background: linear-gradient(135deg, #fff3cd 0%, #ffeaa7 100%);
        color: #856404;
    }

    .status-defaulter {
        background: linear-gradient(135deg, #f8d7da 0%, #f5c6cb 100%);
        color: #721c24;
    }

    .charts-section {
        background: rgba(255, 255, 255, 0.98);
        border-radius: 20px;
        padding: 40px;
        margin-top: 40px;
        box-shadow: 0 20px 60px rgba(0, 0, 0, 0.1);
        border: 1px solid rgba(255, 255, 255, 0.5);
    }

    .charts-section h2 {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        background-clip: text;
        margin-bottom: 35px;
        font-size: 1.9em;
        font-weight: 800;
    }

    .charts-grid {
        display: grid;
        grid-template-columns: repeat(auto-fit, minmax(420px, 1fr));
        gap: 35px;
    }

    .chart-container {
        background: linear-gradient(135deg, #f8faff 0%, #f0f4ff 100%);
        border-radius: 15px;
        padding: 25px;
        box-shadow: 0 10px 35px rgba(0, 0, 0, 0.08);
        border: 2px solid rgba(102, 126, 234, 0.1);
        transition: all 0.3s ease;
    }

    .chart-container:hover {
        box-shadow: 0 20px 50px rgba(102, 126, 234, 0.2);
        transform: translateY(-5px);
    }

    .chart-title {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        background-clip: text;
        font-weight: 800;
        margin-bottom: 20px;
        font-size: 1.2em;
    }

    .admin-section {
        background: rgba(255, 255, 255, 0.98);
        border-radius: 20px;
        padding: 35px;
        margin-top: 40px;
        box-shadow: 0 20px 60px rgba(0, 0, 0, 0.1);
        border-left: 5px solid #dc3545;
        border: 1px solid rgba(255, 255, 255, 0.5);
    }

    .admin-section h2 {
        color: #dc3545;
        margin-bottom: 25px;
        font-size: 1.9em;
        font-weight: 800;
    }

    .admin-info {
        padding: 20px;
        background: linear-gradient(135deg, #fff5f5 0%, #ffe0e0 100%);
        border-radius: 12px;
        margin-bottom: 20px;
        border-left: 4px solid #dc3545;
    }

    .admin-info p {
        color: #555;
        margin-bottom: 10px;
        font-weight: 600;
    }

    .admin-info small {
        color: #999;
    }

    table {
        width: 100%;
        border-collapse: collapse;
        margin-bottom: 30px;
    }

    th {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        padding: 18px;
        text-align: left;
        font-weight: 700;
        font-size: 1.05em;
    }

    td {
        padding: 15px 18px;
        border-bottom: 2px solid #f0f0f0;
    }

    tr:hover {
        background: linear-gradient(135deg, #f8faff 0%, #f0f4ff 100%);
    }

    .safe { background: linear-gradient(135deg, #d4edda 0%, #c3e6cb 100%); color: #155724; }
    .warning { background: linear-gradient(135deg, #fff3cd 0%, #ffeaa7 100%); color: #856404; }
    .defaulter { background: linear-gradient(135deg, #f8d7da 0%, #f5c6cb 100%); color: #721c24; }

    .stats-summary {
        display: grid;
        grid-template-columns: repeat(auto-fit, minmax(220px, 1fr));
        gap: 25px;
        margin-bottom: 40px;
    }

    .stat-box {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        padding: 30px;
        border-radius: 15px;
        text-align: center;
        box-shadow: 0 15px 40px rgba(102, 126, 234, 0.3);
        transition: all 0.3s ease;
        border: 1px solid rgba(255, 255, 255, 0.2);
    }

    .stat-box:hover {
        transform: translateY(-8px);
        box-shadow: 0 25px 60px rgba(102, 126, 234, 0.4);
    }

    .stat-box h3 {
        font-size: 0.95em;
        opacity: 0.95;
        margin-bottom: 12px;
        font-weight: 700;
    }

    .stat-box .value {
        font-size: 2.5em;
        font-weight: 900;
        letter-spacing: 2px;
    }

    footer {
        text-align: center;
        color: rgba(255, 255, 255, 0.9);
        margin-top: 50px;
        padding: 30px;
        font-weight: 600;
        font-size: 1.05em;
    }

    .flash-message {
        padding: 18px;
        border-radius: 12px;
        margin-bottom: 25px;
        font-weight: 600;
        border-left: 5px solid;
        animation: slideDown 0.4s ease;
    }

    @keyframes slideDown {
        from {
            opacity: 0;
            transform: translateY(-20px);
        }
        to {
            opacity: 1;
            transform: translateY(0);
        }
    }

    .flash-success {
        background: linear-gradient(135deg, #d4edda 0%, #c3e6cb 100%);
        color: #155724;
        border-color: #28a745;
    }

    .flash-warning {
        background: linear-gradient(135deg, #fff3cd 0%, #ffeaa7 100%);
        color: #856404;
        border-color: #ffc107;
    }

    .import-section {
        background: rgba(255, 255, 255, 0.98);
        border-radius: 20px;
        padding: 35px;
        margin-bottom: 25px;
        box-shadow: 0 20px 60px rgba(0, 0, 0, 0.1);
        border: 1px solid rgba(255, 255, 255, 0.5);
    }

    .import-section h3 {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        background-clip: text;
        margin-bottom: 20px;
        font-weight: 800;
    }

    .file-input-wrapper {
        position: relative;
        display: inline-block;
        cursor: pointer;
    }

    .file-input-wrapper input[type="file"] {
        position: absolute;
        left: -9999px;
    }

    .file-input-label {
        display: inline-flex;
        align-items: center;
        gap: 10px;
        padding: 12px 20px;
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        border-radius: 10px;
        font-weight: 700;
        cursor: pointer;
        transition: all 0.3s ease;
    }

    .file-input-label:hover {
        transform: translateY(-2px);
        box-shadow: 0 10px 30px rgba(102, 126, 234, 0.4);
    }

    .file-name {
        color: #667eea;
        font-weight: 600;
        margin-left: 10px;
    }

    .import-note {
        background: linear-gradient(135deg, #fff3cd 0%, #ffeaa7 100%);
        border-left: 4px solid #ffc107;
        padding: 15px;
        border-radius: 8px;
        margin-top: 15px;
        font-size: 0.9em;
        color: #856404;
    }

    .date-selector {
        background: rgba(255, 255, 255, 0.98);
        border-radius: 18px;
        padding: 25px;
        margin-bottom: 25px;
        box-shadow: 0 15px 50px rgba(0, 0, 0, 0.1);
        display: flex;
        gap: 20px;
        align-items: center;
        flex-wrap: wrap;
        border: 1px solid rgba(255, 255, 255, 0.5);
    }

    .date-selector label {
        font-weight: 700;
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        background-clip: text;
        font-size: 1.1em;
    }

    .date-selector select {
        padding: 12px 16px;
        border: 2px solid #e0e0e0;
        border-radius: 10px;
        font-size: 1em;
        font-weight: 600;
        transition: all 0.3s ease;
    }

    .date-selector select:focus {
        outline: none;
        border-color: #667eea;
        box-shadow: 0 0 15px rgba(102, 126, 234, 0.3);
    }

    .date-label {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        padding: 10px 20px;
        border-radius: 20px;
        font-weight: 700;
        font-size: 1.05em;
    }

    .attendance-form {
        background: rgba(255, 255, 255, 0.98);
        border-radius: 20px;
        padding: 35px;
        box-shadow: 0 20px 60px rgba(0, 0, 0, 0.1);
        border: 1px solid rgba(255, 255, 255, 0.5);
    }

    .attendance-table {
        width: 100%;
        border-collapse: collapse;
    }

    .attendance-table thead {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
    }

    .attendance-table th {
        color: white;
        padding: 18px;
        text-align: left;
        font-weight: 700;
        font-size: 1.05em;
    }

    .attendance-table td {
        padding: 15px 18px;
        border-bottom: 2px solid #f0f0f0;
    }

    .attendance-table tr:hover {
        background: linear-gradient(135deg, #f8faff 0%, #f0f4ff 100%);
    }

    .attendance-table select {
        padding: 10px 14px;
        border: 2px solid #e0e0e0;
        border-radius: 8px;
        font-size: 0.9em;
        font-weight: 600;
        transition: all 0.3s ease;
    }

    .attendance-table select:focus {
        outline: none;
        border-color: #667eea;
        box-shadow: 0 0 10px rgba(102, 126, 234, 0.3);
    }

    .attendance-table input {
        padding: 10px 14px;
        border: 2px solid #e0e0e0;
        border-radius: 8px;
        font-size: 0.9em;
        width: 100%;
        transition: all 0.3s ease;
    }

    .attendance-table input:focus {
        outline: none;
        border-color: #667eea;
        box-shadow: 0 0 10px rgba(102, 126, 234, 0.3);
    }

    @media (max-width: 768px) {
        .header {
            flex-direction: column;
            text-align: center;
            padding: 30px;
        }

        .header-left {
            flex-direction: column;
            width: 100%;
        }

        .header-title h1 {
            font-size: 2em;
        }

        .classes-grid {
            grid-template-columns: 1fr;
        }

        .charts-grid {
            grid-template-columns: 1fr;
        }

        .form-group {
            flex-direction: column;
        }

        .form-group input {
            width: 100%;
        }

        table {
            font-size: 0.9em;
        }

        td, th {
            padding: 10px;
        }

        .stat-box .value {
            font-size: 2em;
        }
    }
    .status-btn {
    padding: 8px 14px;
    border-radius: 10px;
    border: 2px solid #ddd;
    background: #f8f9fa;
    cursor: pointer;
    font-weight: 700;
    transition: all 0.25s ease;
}

.status-btn:hover {
    transform: translateY(-2px);
}

.active-present {
    background: linear-gradient(135deg, #28a745, #20c997);
    color: white;
    border-color: #28a745;
}

.active-absent {
    background: linear-gradient(135deg, #dc3545, #c82333);
    color: white;
    border-color: #dc3545;
}

.active-od {
    background: linear-gradient(135deg, #ffc107, #ff9800);
    color: white;
    border-color: #ffc107;
}
/* ===== STATUS BUTTONS ===== */
.status-group {
    display: flex;
    gap: 12px;
}

.status-btn {
    position: relative;
}

.status-btn input {
    display: none;
}

.status-btn label {
    padding: 10px 18px;
    border-radius: 20px;
    font-weight: 700;
    cursor: pointer;
    border: 2px solid #ddd;
    transition: all 0.25s ease;
    user-select: none;
}

/* COLORS */
.status-present label { color: #28a745; }
.status-absent label { color: #dc3545; }
.status-od label { color: #ffc107; }

/* ACTIVE STATES */
.status-present input:checked + label {
    background: #28a745;
    color: white;
    animation: blinkGreen 0.4s;
}

.status-absent input:checked + label {
    background: #dc3545;
    color: white;
    animation: blinkRed 0.4s;
}

.status-od input:checked + label {
    background: #ffc107;
    color: #333;
    animation: blinkYellow 0.4s;
}

/* BLINK ANIMATIONS */
@keyframes blinkGreen {
    0% { box-shadow: 0 0 0 0 rgba(40,167,69,0.9); }
    100% { box-shadow: 0 0 0 10px rgba(40,167,69,0); }
}

@keyframes blinkRed {
    0% { box-shadow: 0 0 0 0 rgba(220,53,69,0.9); }
    100% { box-shadow: 0 0 0 10px rgba(220,53,69,0); }
}

@keyframes blinkYellow {
    0% { box-shadow: 0 0 0 0 rgba(255,193,7,0.9); }
    100% { box-shadow: 0 0 0 10px rgba(255,193,7,0); }
}
.charts-grid {
    display: grid;
    gap: 30px;
}

@media (min-width: 900px) {
    .charts-grid {
        grid-template-columns: 1fr 1fr;
    }
}



</style>
"""

# ====================== ROUTES ======================

@app.route("/admin_logout")
def admin_logout():
    session.pop("admin", None)
    flash("Admin logged out 🔒", "success")
    return redirect(request.referrer or "/")
@app.route("/admin_reset", methods=["POST"])
def admin_reset():
    if not session.get("admin"):
        flash("❌ Admin access required", "warning")
        return redirect("/")

    db = get_db()
    db.execute("DELETE FROM admin")
    db.commit()
    db.close()

    session.pop("admin", None)

    flash("🧹 Admin password reset successfully. Please set a new password.", "success")
    return redirect("/")


@app.route("/admin_unlock", methods=["POST"])
def admin_unlock():
    password = request.form.get("password")
    confirm = request.form.get("confirm")

    db = get_db()
    existing = db.execute("SELECT password FROM admin").fetchone()

    # FIRST TIME SETUP
    if not existing:
        if not password or not confirm:
            flash("❌ Please fill both password fields", "warning")
            return redirect("/")

        if password != confirm:
            flash("❌ Passwords do not match", "warning")
            return redirect("/")

        if len(password) < 6:
            flash("❌ Password must be at least 6 characters", "warning")
            return redirect("/")

        hashed = generate_password_hash(password)
        db.execute("INSERT INTO admin (password) VALUES (?)", (hashed,))
        db.commit()
        db.close()

        session["admin"] = True
        flash("✅ Admin password set successfully!", "success")
        return redirect("/")

    # NORMAL LOGIN
    stored_hash = existing[0]

    if not check_password_hash(stored_hash, password):
        db.close()
        flash("❌ Incorrect admin password", "warning")
        return redirect("/")

    db.close()
    session["admin"] = True
    flash("🔐 Admin login successful", "success")
    return redirect("/")



@app.route("/")
def dashboard():
    db = get_db()
    admin_exists = bool(
    db.execute("SELECT password FROM admin").fetchone()
)

    classes_db = db.execute("SELECT id, class_name FROM classes").fetchall()

    class_stats = []
    for cid, cname in classes_db:
        students = db.execute("SELECT id FROM students WHERE class_id=?", (cid,)).fetchall()
        total_students = len(students)

        safe = warning = defaulters = 0
        total_percent = 0

        for (sid,) in students:
            stats = get_student_attendance_stats(sid, db)
            total_percent += stats["percent"]

            if stats["status"] == "Safe":
                safe += 1
            elif stats["status"] == "Warning":
                warning += 1
            else:
                defaulters += 1

        avg_percent = round(total_percent / total_students, 2) if total_students else 0

        class_stats.append({
            "id": cid,
            "name": cname,
            "total_students": total_students,
            "safe": safe,
            "warning": warning,
            "defaulters": defaulters,
            "avg_percent": avg_percent
        })

    db.close()

    html = """
    <!DOCTYPE html>
    <html lang="en">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>Attendance Monitor - Dashboard</title>
        <script src="https://cdn.jsdelivr.net/npm/chart.js"></script>
        """ + BASE_STYLES + """
    </head>
    <body>
        <div class="container">
            <div class="header">
                <div class="header-left">
                    <div class="header-logo">
                        📊
                    </div>
                    <div class="header-title">
                        <h1> Perfect Attendance Monitor</h1>
                        <p>Smart Class Attendance Monitor System</p>
                    </div>
                </div>
                <div class="header-actions">
                    {% if session.get('admin') %}
                        <span style="color: #667eea; font-weight: 700; font-size: 1.1em;">🔐 Admin Mode</span>
                        <a href="/admin_logout" class="btn btn-danger">🔓 Logout</a>
                    {% else %}
                        <button class="btn btn-primary" onclick="showAdminModal()">🔒 Admin Login</button>
                    {% endif %}
                   {% if classes %}
<div class="charts-section" style="margin-top: 30px;">
    <h2 style="text-align:center;">📊 Overall Class Analytics</h2>

    <div class="charts-grid"
         style="grid-template-columns: repeat(auto-fit, minmax(380px, 1fr));">
         
        <div class="chart-container">
            <div class="chart-title">📚 Average Attendance %</div>
            <canvas id="avgAttendanceChart"></canvas>
        </div>

        <div class="chart-container">
            <div class="chart-title">🚦 Student Status Distribution</div>
            <canvas id="statusDistributionChart"></canvas>
        </div>
    </div>
</div>
{% endif %}


                </div>
            </div>

            {% with messages = get_flashed_messages(with_categories=true) %}
                {% if messages %}
                    {% for category, message in messages %}
                        <div class="flash-message flash-{{ 'success' if category == 'success' else 'warning' }}">
                            {{ message }}
                        </div>
                    {% endfor %}
                {% endif %}
            {% endwith %}

            <div class="add-class-form">
                <h2>➕ Add New Class</h2>
                <form method="POST" action="/add_class">
                    <div class="form-group">
                        <input type="text" name="class_name" placeholder="Enter class name (e.g., Class AIML)" required>
                        <button type="submit" class="btn btn-primary">Add Class</button>
                    </div>
                </form>
            </div>

            {% if classes %}
                <div class="classes-grid">
                    {% for class in classes %}
                        <div class="class-card">
                            <div class="class-header">
                                <span class="class-name">📚 {{ class.name }}</span>
                                <span class="class-badge">{{ class.total_students }} Students</span>
                            </div>

                            <div class="class-stats">
                                <div class="stat-item">
                                    <span>👥 Total Students</span>
                                    <span class="stat-value">{{ class.total_students }}</span>
                                </div>
                                <div class="stat-item">
                                    <span>🟢 Safe (≥75%)</span>
                                    <span class="stat-value">{{ class.safe }}</span>
                                </div>
                                <div class="stat-item">
                                    <span>🟡 Warning (65-74%)</span>
                                    <span class="stat-value">{{ class.warning }}</span>
                                </div>
                                <div class="stat-item">
                                    <span>🔴 Defaulters (<65%)</span>
                                    <span class="stat-value stat-defaulter">{{ class.defaulters }}</span>
                                </div>
                                <div class="stat-item">
                                    <span>📊 Avg Attendance</span>
                                    <span class="stat-value">{{ class.avg_percent }}%</span>
                                </div>
                            </div>

                            <div class="class-actions">
                                <a href="/attendance/{{ class.id }}" class="btn btn-primary">📋 Mark Attendance</a>
                                <a href="/report/total/{{ class.id }}" class="btn btn-secondary">📊 View Report</a>
                            </div>

                            {% if session.get('admin') %}
                                <div class="class-actions" style="margin-top: 15px;">
                                    <form method="POST" action="/delete_class/{{ class.id }}" 
                                          onsubmit="return confirm('⚠️ This will delete the entire class and all records! Are you sure?');">
                                        <button type="submit" class="btn btn-danger" style="width: 100%;">🗑️ Delete Class</button>
                                    </form>
                                </div>
                            {% endif %}
                        </div>
                    {% endfor %}
                </div>
            {% else %}
                <div style="background: rgba(255, 255, 255, 0.95); border-radius: 20px; padding: 60px; text-align: center; box-shadow: 0 20px 60px rgba(0, 0, 0, 0.1);">
                    <div style="font-size: 4em; margin-bottom: 20px;">📚</div>
                    <h2 style="color: #667eea; margin-bottom: 15px; font-size: 1.8em;">No Classes Yet</h2>
                    <p style="color: #666; font-size: 1.1em; margin-bottom: 30px;">Add your first class to get started!</p>
                    <button class="btn btn-primary" onclick="document.querySelector('input[name=class_name]').focus()">Create First Class</button>
                </div>
            {% endif %}

           {% if session.get('admin') %}
<div class="admin-section">
    <h2>🔐 Admin Control Panel</h2>

    <div class="admin-info">
        <p>✅ <strong>Admin Status:</strong> Logged In</p>
        <small>You have full access to all features.</small>
    </div>

    <div style="display:flex; gap:15px; flex-wrap:wrap;">
        <a href="/admin_logout" class="btn btn-secondary">
            🔓 Logout
        </a>

        <form method="POST" action="/admin_reset"
              onsubmit="return confirm('⚠️ This will RESET admin password. Continue?');">
            <button type="submit" class="btn btn-danger">
                🧹 Reset Admin Password
            </button>
        </form>
    </div>
</div>
{% endif %}

        <div id="adminModal" style="display: none; position: fixed; top: 0; left: 0; 
                                    width: 100%; height: 100%; background: rgba(0,0,0,0.8); 
                                    z-index: 9999; align-items: center; justify-content: center;">
            <div style="background: white; padding: 50px; border-radius: 20px; 
                        box-shadow: 0 30px 90px rgba(0,0,0,0.4); max-width: 450px; width: 95%; animation: slideUp 0.4s ease;">
                <h2 style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); 
                           -webkit-background-clip: text; -webkit-text-fill-color: transparent;
                           background-clip: text;
                           margin-bottom: 25px; font-size: 1.8em; font-weight: 800;">🔒 Admin Login</h2>
                <form method="POST" action="/admin_unlock">
                   

    <div style="margin-bottom: 18px;">
        <input type="password" name="password"
               placeholder="Enter admin password"
               required
               style="width:100%; padding:15px; border-radius:12px; border:2px solid #ddd;">
    </div>

    {% if not admin_exists %}
    <div style="margin-bottom: 18px;">
        <input type="password" name="confirm"
               placeholder="Confirm admin password"
               required
               style="width:100%; padding:15px; border-radius:12px; border:2px solid #ddd;">
    </div>
    {% endif %}

    <div style="display:flex; gap:12px;">
        <button type="submit" class="btn btn-primary" style="flex:1;">
            {% if admin_exists %}Login{% else %}Set Password{% endif %}
        </button>

        <button type="button" class="btn btn-secondary" style="flex:1;"
                onclick="document.getElementById('adminModal').style.display='none'">
            Cancel
        </button>
    </div>
</form>

                
            </div>
        </div>

        <footer>
            <p>🎓 Perfect Attendance Monitor | © 2026 </p>
            <p>Developed by R RAHUL | VIJEY ABINESSH | [AIML DEPT] </p>
            <P>ASTHRA TECH PVT LTD</P>
            <P>Contact: 8939650849 | 99419 33332 </P>
            <P>CHENNAI INSTITUTE OF TECHNOLOGY</P>
        </footer>

        <script>
            function showAdminModal() {
                document.getElementById('adminModal').style.display = 'flex';
            }

            document.getElementById('adminModal').addEventListener('click', function(e) {
                if (e.target === this) {
                    this.style.display = 'none';
                }
            });
        </script>
        <script>
    const classStats = {{ classes | tojson }};

    const labels = classStats.map(c => c.name);
    const avgPercent = classStats.map(c => c.avg_percent);

    const safe = classStats.reduce((a,c)=>a+c.safe,0);
    const warning = classStats.reduce((a,c)=>a+c.warning,0);
    const defaulters = classStats.reduce((a,c)=>a+c.defaulters,0);

    new Chart(document.getElementById("avgAttendanceChart"), {
        type: "bar",
        data: {
            labels: labels,
            datasets: [{
                label: "Avg Attendance %",
                data: avgPercent,
                backgroundColor: "#667eea",
                borderRadius: 8
            }]
        },
        options: {
            scales: {
                y: {
                    beginAtZero: true,
                    max: 100,
                    ticks: { callback: v => v + "%" }
                }
            }
        }
    });

    new Chart(document.getElementById("statusDistributionChart"), {
        type: "doughnut",
        data: {
            labels: ["🟢 Safe", "🟡 Warning", "🔴 Defaulter"],
            datasets: [{
                data: [safe, warning, defaulters],
                backgroundColor: ["#28a745", "#ffc107", "#dc3545"]
            }]
        }
    });
</script>

    </body>
    </html>
    """

    return render_template_string(
    html,
    classes=class_stats,
    session=session,
    admin_exists=admin_exists,
    get_flashed_messages=get_flashed_messages
)



@app.route("/add_class", methods=["POST"])
def add_class():
    name = request.form.get("class_name", "").strip()
    
    if not name:
        flash("❌ Class name cannot be empty!", "warning")
        return redirect("/")
    
    try:
        db = get_db()
        c = db.cursor()
        
        # Check if class already exists
        existing = c.execute("SELECT id FROM classes WHERE class_name=?", (name,)).fetchone()
        if existing:
            flash(f"⚠️ Class '{name}' already exists!", "warning")
            db.close()
            return redirect("/")
        
        # Insert the class
        c.execute("INSERT INTO classes (class_name) VALUES (?)", (name,))
        db.commit()
        
        # Verify it was inserted
        verify = c.execute("SELECT id FROM classes WHERE class_name=?", (name,)).fetchone()
        
        if verify:
            flash(f"✅ Class '{name}' created successfully!", "success")
        else:
            flash(f"❌ Failed to create class '{name}'", "warning")
        
        db.close()
    except Exception as e:
        flash(f"❌ Error creating class: {str(e)}", "warning")
    
    return redirect("/")


@app.route("/attendance/<int:class_id>")
def attendance(class_id):
    db = get_db()
    today = date.today().isoformat()

    selected_date = request.args.get("date", today)

    dates_raw = db.execute("""
        SELECT DISTINCT date FROM attendance
        WHERE student_id IN (
            SELECT id FROM students WHERE class_id=?
        )
    """, (class_id,)).fetchall()

    dates = [d[0] for d in dates_raw]

    if today not in dates:
        dates.append(today)

    dates.sort(reverse=True)

    students = db.execute("""
        SELECT * FROM students
        WHERE class_id=?
        ORDER BY roll_no
    """, (class_id,)).fetchall()

    # Get class name
    class_name = db.execute("SELECT class_name FROM classes WHERE id=?", (class_id,)).fetchone()[0]

    attendance_data = db.execute("""
        SELECT student_id, status, od_reason
        FROM attendance
        WHERE date=?
    """, (selected_date,)).fetchall()

    percent_map = {}
    status_map = {}

    if session.get("admin"):
        for s in students:
            sid = s[0]
            stats = get_student_attendance_stats(sid, db)
            percent_map[sid] = stats["percent"]
            status_map[sid] = stats

    db.close()

    attendance_dict = {str(a[0]): a for a in attendance_data}

    html = """
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>Mark Attendance</title>
        """ + BASE_STYLES + """
    </head>
    <body>
        <div class="container">
            <div class="header">
                <h1 style="color: #667eea;">📋 Mark Attendance - {{ class_name }}</h1>
                <a href="/" class="btn btn-secondary">← Back to Dashboard</a>
            </div>

            {% with messages = get_flashed_messages(with_categories=true) %}
                {% if messages %}
                    {% for category, message in messages %}
                        <div class="flash-message flash-{{ 'success' if category == 'success' else 'warning' }}">
                            {{ message }}
                        </div>
                    {% endfor %}
                {% endif %}
            {% endwith %}

            {% if session.get('admin') %}
                <div class="import-section">
                    <h3>📥 Import Students from File</h3>
                    <form method="POST" action="/import/{{ class_id }}" enctype="multipart/form-data">
                        <div style="display: flex; gap: 15px; align-items: center; flex-wrap: wrap;">
                            <div class="file-input-wrapper">
                                <input type="file" id="fileInput" name="file" 
                                       accept=".xlsx,.xls,.docx,.pdf" required
                                       onchange="updateFileName(this)">
                                <label for="fileInput" class="file-input-label">
                                    📁 Choose File
                                </label>
                            </div>
                            <span class="file-name" id="fileName">No file selected</span>
                            <button type="submit" class="btn btn-primary">⬆️ Import</button>
                        </div>
                        <div class="import-note">
                            <strong>✓ Supported formats:</strong> Excel (.xlsx), Word (.docx), PDF (.pdf)<br>
                            <strong>✓ Format:</strong> Excel/Word tables with Roll No & Name columns | PDF: one student per line as "Roll - Name"
                        </div>
                    </form>
                </div>
            {% endif %}
            <div style="
    background: rgba(255,255,255,0.98);
    border-radius: 18px;
    padding: 20px;
    margin-bottom: 20px;
    box-shadow: 0 12px 40px rgba(0,0,0,0.1);
">
    <h3 style="
        margin-bottom: 14px;
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        font-weight: 800;
    ">
        📅 Quick Access – Last 10 Days
    </h3>

    <div style="display: flex; gap: 12px; flex-wrap: wrap;">
        {% for d in dates[:10] %}
            <a href="/attendance/{{ class_id }}?date={{ d }}"
               style="
                   padding: 10px 18px;
                   border-radius: 20px;
                   font-weight: 700;
                   text-decoration: none;
                   transition: all 0.3s ease;
                   {% if d == selected_date %}
                       background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                       color: white;
                       box-shadow: 0 8px 25px rgba(102,126,234,0.5);
                   {% elif d in attendance_dict.values()|map(attribute=0)|list %}
                       background: #f0f4ff;
                       color: #667eea;
                       border: 2px solid #667eea;
                   {% else %}
                       background: #f5f5f5;
                       color: #999;
                       border: 2px dashed #ccc;
                   {% endif %}
               ">
               {{ d }}
            </a>
        {% endfor %}
    </div>
</div>

            
           {% if session.get('admin') %}
<div style="
    background: rgba(255,255,255,0.98);
    border-radius: 18px;
    padding: 25px;
    margin-bottom: 25px;
    box-shadow: 0 12px 40px rgba(0,0,0,0.1);
">
    <h3 style="
        margin-bottom: 18px;
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        font-weight: 800;
    ">
        📤 Export Attendance (Date Range)
    </h3>

    <form method="POST" action="/export/range/excel/{{ class_id }}"
          style="display: flex; gap: 15px; flex-wrap: wrap; align-items: end;">

        <div>
            <label style="font-weight:700;">From</label><br>
            <input type="date" name="start"
                   style="padding:10px 14px; border-radius:10px; border:2px solid #ddd;"
                   required>
        </div>

        <div>
            <label style="font-weight:700;">To</label><br>
            <input type="date" name="end"
                   style="padding:10px 14px; border-radius:10px; border:2px solid #ddd;"
                   required>
        </div>

        <div>
            <button type="submit" class="btn btn-primary">
                📥 Export Excel
            </button>
        </div>
    </form>

    <p style="margin-top:12px; color:#666; font-size:0.9em;">
        🔹 Use this for monthly / custom reports.  
        🔹 Quick daily access is available above.
    </p>
</div>
{% endif %}



            <div class="date-selector">
                <label>📅 Select Date:</label>
                <select onchange="window.location.href = '/attendance/{{ class_id }}?date=' + this.value;">
                    {% for d in dates %}
                        <option value="{{ d }}" {% if d == selected_date %}selected{% endif %}>{{ d }}</option>
                    {% endfor %}
                </select>
                <span class="date-label">{{ selected_date }}</span>
            </div>

            <form method="POST" action="/save_attendance" class="attendance-form">
                <input type="hidden" name="class_id" value="{{ class_id }}">
                <input type="hidden" name="selected_date" value="{{ selected_date }}">

                <table class="attendance-table">
                    <thead>
                        <tr>
                            <th>Roll No</th>
                            <th>Name</th>
                            <th>Status</th>
                            <th>OD Reason</th>
                            {% if session.get('admin') %}<th>Analytics</th>{% endif %}
                        </tr>
                    </thead>
                    <tbody>
                        {% for student in students %}
                            {% set sid = student[0] %}
                            {% set att = attendance_dict.get(sid|string) %}
                            <tr>
                                <td><strong style="font-size: 1.1em;">{{ student[1] }}</strong></td>
                                <td style="font-weight: 600;">
    {{ student[2] }}
    {% if session.get('admin') %}
        <br>
        <a href="/report/student/{{ sid }}"
           class="btn btn-secondary"
           style="margin-top:6px; padding:6px 10px; font-size:0.8em;">
           📊 Analytics
        </a>
    {% endif %}
</td>
                                <td>
                                    
                                       <div class="status-group">

    <div class="status-btn status-present">
        <input type="radio" name="status_{{ sid }}" value="Present"
               id="p{{ sid }}"
               {% if att and att[1]=='Present' %}checked{% endif %}>
        <label for="p{{ sid }}">PRESENT</label>
    </div>

    <div class="status-btn status-absent">
        <input type="radio" name="status_{{ sid }}" value="Absent"
               id="a{{ sid }}"
               {% if att and att[1]=='Absent' %}checked{% endif %}>
        <label for="a{{ sid }}">ABSENT</label>
    </div>

    <div class="status-btn status-od">
        <input type="radio" name="status_{{ sid }}" value="OD"
               id="o{{ sid }}"
               {% if att and att[1]=='OD' %}checked{% endif %}>
        <label for="o{{ sid }}">OD</label>
    </div>

</div>


                                </td>
                                <td>
                                    <input type="text" name="r_{{ sid }}" placeholder="Reason" 
                                           value="{% if att and att[2] %}{{ att[2] }}{% endif %}">
                                </td>
                                {% if session.get('admin') %}
                                    <td>
                                        {% set stats = status_map.get(sid) %}
                                        {% if stats %}
                                            <span class="status-badge status-{{ stats['class'] }}" style="white-space: nowrap;">
                                                {{ stats['badge'] }} {{ stats['status'] }}<br><small>{{ percent_map.get(sid, 0) }}%</small>
                                            </span>
                                        {% else %}
                                            <span style="color: #999;">No data</span>
                                        {% endif %}
                                    </td>
                                {% endif %}
                            </tr>
                        {% endfor %}
                    </tbody>
                </table>

                <div style="margin-top: 25px; display: flex; gap: 12px;">
                    <button type="submit" class="btn btn-primary">💾 Save Attendance</button>
                    <a href="/" class="btn btn-secondary">← Cancel</a>
                </div>
            </form>
        </div>
        {% if session.get('admin') %}
<div style="margin-bottom: 20px; display:flex; gap:10px;">
    <a href="/report/day/{{ class_id }}?date={{ selected_date }}"
       class="btn btn-primary">📅 Day Report</a>

    <a href="/export/day/excel/{{ class_id }}?date={{ selected_date }}"
       class="btn btn-secondary">📥 Excel</a>
</div>
{% endif %}


        <script>
            function updateFileName(input) {
                const fileName = document.getElementById('fileName');
                if (input.files && input.files[0]) {
                    fileName.textContent = '✓ ' + input.files[0].name;
                    fileName.style.color = '#28a745';
                } else {
                    fileName.textContent = 'No file selected';
                    fileName.style.color = '#667eea';
                }
            }
        </script>
        

    </body>
    </html>
    """

    return render_template_string(html, 
                                 class_id=class_id, 
                                 class_name=class_name,
                                 dates=dates, 
                                 selected_date=selected_date,
                                 students=students,
                                 attendance_dict=attendance_dict,
                                 percent_map=percent_map,
                                 status_map=status_map,
                                 session=session,
                                 get_flashed_messages=get_flashed_messages)

@app.route("/export/range/excel/<int:class_id>", methods=["POST"])
def export_range_excel(class_id):
    if not session.get("admin"):
        return "Admin access required"

    start = request.form.get("start")
    end = request.form.get("end")

    if not start or not end:
        flash("❌ Please select both From and To dates", "warning")
        return redirect(f"/attendance/{class_id}")

    db = get_db()
    rows = db.execute("""
        SELECT s.roll_no, s.name, a.date, a.status, a.od_reason
        FROM attendance a
        JOIN students s ON a.student_id = s.id
        WHERE s.class_id=? AND a.date BETWEEN ? AND ?
        ORDER BY a.date, s.roll_no
    """, (class_id, start, end)).fetchall()

    db.close()

    if not rows:
        flash("⚠️ No attendance records found in selected range", "warning")
        return redirect(f"/attendance/{class_id}")

    wb = Workbook()
    ws = wb.active
    ws.title = "Attendance Range Report"

    ws.append(["Roll No", "Name", "Date", "Status", "OD Reason"])
    for r in rows:
        ws.append(r)

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx")
    wb.save(tmp.name)

    return send_file(
        tmp.name,
        as_attachment=True,
        download_name=f"Attendance_{start}_to_{end}.xlsx"
    )

@app.route("/delete_class/<int:class_id>", methods=["POST"])
def delete_class(class_id):
    if not session.get("admin"):
        flash("❌ Admin access required", "warning")
        return redirect("/")

    db = get_db()
    class_name = db.execute("SELECT class_name FROM classes WHERE id=?", (class_id,)).fetchone()
    
    if not class_name:
        db.close()
        flash("❌ Class not found", "warning")
        return redirect("/")
    
    class_name = class_name[0]
    
    try:
        db.execute("""
            DELETE FROM attendance
            WHERE student_id IN (
                SELECT id FROM students WHERE class_id=?
            )
        """, (class_id,))

        db.execute("DELETE FROM students WHERE class_id=?", (class_id,))
        db.execute("DELETE FROM classes WHERE id=?", (class_id,))
        db.commit()
        flash(f"🗑️ Class '{class_name}' deleted successfully!", "success")
    except Exception as e:
        flash(f"❌ Error deleting class: {str(e)}", "warning")
    finally:
        db.close()
    
    return redirect("/")


@app.route("/report/day/<int:class_id>")
def day_report(class_id):
    if not session.get("admin"):
        return "Admin access required"

    report_date = request.args.get("date", date.today().isoformat())
    db = get_db()

    rows = db.execute("""
        SELECT s.roll_no, s.name, a.status, a.od_reason
        FROM attendance a
        JOIN students s ON a.student_id = s.id
        WHERE s.class_id=? AND a.date=?
        ORDER BY s.roll_no
    """, (class_id, report_date)).fetchall()

    db.close()

    present = [r for r in rows if r[2] == "Present"]
    absent = [r for r in rows if r[2] == "Absent"]
    od = [r for r in rows if r[2] == "OD"]

    html = """
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <title>Day Report</title>
        """ + BASE_STYLES + """
    </head>
    <body>
        <div class="container">
            <div class="header">
                <h1 style="color: #667eea;">📊 Day Attendance Report</h1>
                <div>
                    <a href="/attendance/{{ class_id }}" class="btn btn-secondary">← Back</a>
                </div>
            </div>

            <div style="background: rgba(255, 255, 255, 0.98); border-radius: 20px; padding: 35px; box-shadow: 0 20px 60px rgba(0, 0, 0, 0.1); border: 1px solid rgba(255, 255, 255, 0.5);">
                <h2 style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); -webkit-background-clip: text; -webkit-text-fill-color: transparent; background-clip: text; margin-bottom: 30px; font-size: 1.8em;">📅 {{ report_date }}</h2>

                <div class="stats-summary">
                    <div class="stat-box" style="background: linear-gradient(135deg, #28a745 0%, #20c997 100%);">
                        <h3>✅ Present</h3>
                        <div class="value">{{ present|length }}</div>
                    </div>
                    <div class="stat-box" style="background: linear-gradient(135deg, #dc3545 0%, #c82333 100%);">
                        <h3>❌ Absent</h3>
                        <div class="value">{{ absent|length }}</div>
                    </div>
                    <div class="stat-box" style="background: linear-gradient(135deg, #ffc107 0%, #ff9800 100%);">
                        <h3>📝 OD</h3>
                        <div class="value">{{ od|length }}</div>
                    </div>
                </div>

                <h3 style="color: #28a745; margin: 30px 0 15px 0; font-size: 1.5em; font-weight: 800;">✅ Present ({{ present|length }})</h3>
                <table style="width: 100%; border-collapse: collapse; margin-bottom: 30px;">
                    <thead>
                        <tr style="background: linear-gradient(135deg, #28a745 0%, #20c997 100%); color: white;">
                            <th style="padding: 15px; text-align: left; font-weight: 700;">Roll No</th>
                            <th style="padding: 15px; text-align: left; font-weight: 700;">Name</th>
                        </tr>
                    </thead>
                    <tbody>
                        {% for p in present %}
                            <tr style="border-bottom: 2px solid #f0f0f0;">
                                <td style="padding: 12px 15px; font-weight: 600;">{{ p[0] }}</td>
                                <td style="padding: 12px 15px;">{{ p[1] }}</td>
                            </tr>
                        {% else %}
                            <tr><td colspan="2" style="padding: 20px; text-align: center; color: #999;">No records</td></tr>
                        {% endfor %}
                    </tbody>
                </table>

                <h3 style="color: #dc3545; margin: 30px 0 15px 0; font-size: 1.5em; font-weight: 800;">❌ Absent ({{ absent|length }})</h3>
                <table style="width: 100%; border-collapse: collapse; margin-bottom: 30px;">
                    <thead>
                        <tr style="background: linear-gradient(135deg, #dc3545 0%, #c82333 100%); color: white;">
                            <th style="padding: 15px; text-align: left; font-weight: 700;">Roll No</th>
                            <th style="padding: 15px; text-align: left; font-weight: 700;">Name</th>
                        </tr>
                    </thead>
                    <tbody>
                        {% for a in absent %}
                            <tr style="border-bottom: 2px solid #f0f0f0;">
                                <td style="padding: 12px 15px; font-weight: 600;">{{ a[0] }}</td>
                                <td style="padding: 12px 15px;">{{ a[1] }}</td>
                            </tr>
                        {% else %}
                            <tr><td colspan="2" style="padding: 20px; text-align: center; color: #999;">No records</td></tr>
                        {% endfor %}
                    </tbody>
                </table>

                <h3 style="color: #ffc107; margin: 30px 0 15px 0; font-size: 1.5em; font-weight: 800;">📝 OD ({{ od|length }})</h3>
                <table style="width: 100%; border-collapse: collapse;">
                    <thead>
                        <tr style="background: linear-gradient(135deg, #ffc107 0%, #ff9800 100%); color: white;">
                            <th style="padding: 15px; text-align: left; font-weight: 700;">Roll No</th>
                            <th style="padding: 15px; text-align: left; font-weight: 700;">Name</th>
                            <th style="padding: 15px; text-align: left; font-weight: 700;">Reason</th>
                        </tr>
                    </thead>
                    <tbody>
                        {% for o in od %}
                            <tr style="border-bottom: 2px solid #f0f0f0;">
                                <td style="padding: 12px 15px; font-weight: 600;">{{ o[0] }}</td>
                                <td style="padding: 12px 15px;">{{ o[1] }}</td>
                                <td style="padding: 12px 15px; color: #666;">{{ o[3] or '-' }}</td>
                            </tr>
                        {% else %}
                            <tr><td colspan="3" style="padding: 20px; text-align: center; color: #999;">No records</td></tr>
                        {% endfor %}
                    </tbody>
                </table>
            </div>
        </div>
    </body>
    </html>
    """

    return render_template_string(html, 
                                 class_id=class_id,
                                 report_date=report_date,
                                 present=present,
                                 absent=absent,
                                 od=od)


@app.route("/report/total/<int:class_id>")
def total_report(class_id):
    if not session.get("admin"):
        return "Admin access required"

    db = get_db()

    students = db.execute("""
        SELECT id, roll_no, name
        FROM students
        WHERE class_id=?
        ORDER BY roll_no
    """, (class_id,)).fetchall()

    report = []

    for sid, roll, name in students:
        stats = get_student_attendance_stats(sid, db)
        report.append({
            "roll": roll,
            "name": name,
            "total": stats["total"],
            "present": stats["present"],
            "absent": stats["absent"],
            "percent": stats["percent"],
            "status": stats["status"],
            "badge": stats["badge"],
            "color": stats["color"]
        })

    db.close()

    safe_count = sum(1 for r in report if r['status'] == 'Safe')
    warning_count = sum(1 for r in report if r['status'] == 'Warning')
    defaulter_count = sum(1 for r in report if r['status'] == 'Defaulter')

    html = """
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>Total Attendance Report</title>
        <script src="https://cdn.jsdelivr.net/npm/chart.js@3.9.1/dist/chart.min.js"></script>
        """ + BASE_STYLES + """
    </head>
    <body>
        <div class="container">
            <div class="header">
                <h1 style="color: #667eea;">📊 Total Class Attendance Report</h1>
                <div style="display: flex; gap: 12px;">
                    <a href="/export/total/excel/{{ class_id }}" class="btn btn-primary">📥 Excel</a>
                    <a href="/attendance/{{ class_id }}" class="btn btn-secondary">← Back</a>
                </div>
            </div>

            <div style="background: rgba(255, 255, 255, 0.98); border-radius: 20px; padding: 40px; box-shadow: 0 20px 60px rgba(0, 0, 0, 0.1); border: 1px solid rgba(255, 255, 255, 0.5);">
                <div class="stats-summary">
                    <div class="stat-box">
                        <h3>👥 Total Students</h3>
                        <div class="value">{{ report|length }}</div>
                    </div>
                    <div class="stat-box" style="background: linear-gradient(135deg, #28a745 0%, #20c997 100%);">
                        <h3>🟢 Safe (≥75%)</h3>
                        <div class="value">{{ safe_count }}</div>
                    </div>
                    <div class="stat-box" style="background: linear-gradient(135deg, #ffc107 0%, #ff9800 100%);">
                        <h3>🟡 Warning (65-74%)</h3>
                        <div class="value">{{ warning_count }}</div>
                    </div>
                    <div class="stat-box" style="background: linear-gradient(135deg, #dc3545 0%, #c82333 100%);">
                        <h3>🔴 Defaulter (<65%)</h3>
                        <div class="value">{{ defaulter_count }}</div>
                    </div>
                </div>

                <h2 style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); -webkit-background-clip: text; -webkit-text-fill-color: transparent; background-clip: text; margin: 40px 0 25px 0; font-size: 1.8em; font-weight: 800;">📋 Student Details</h2>

                <div style="overflow-x: auto;">
                    <table style="width: 100%; border-collapse: collapse;">
                        <thead>
                            <tr style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); color: white;">
                                <th style="padding: 18px; text-align: left;">Roll No</th>
                                <th style="padding: 18px; text-align: left;">Name</th>
                                <th style="padding: 18px; text-align: center;">Total</th>
                                <th style="padding: 18px; text-align: center;">Present</th>
                                <th style="padding: 18px; text-align: center;">Absent</th>
                                <th style="padding: 18px; text-align: center;">%</th>
                                <th style="padding: 18px; text-align: center;">Status</th>
                            </tr>
                        </thead>
                        <tbody>
                            {% for item in report %}
                                <tr style="border-bottom: 2px solid #f0f0f0;">
                                    <td style="padding: 15px 18px; font-weight: 700;">{{ item.roll }}</td>
                                    <td style="padding: 15px 18px; font-weight: 600;">{{ item.name }}</td>
                                    <td style="padding: 15px 18px; text-align: center; font-weight: 600;">{{ item.total }}</td>
                                    <td style="padding: 15px 18px; text-align: center; color: #28a745; font-weight: 700;">{{ item.present }}</td>
                                    <td style="padding: 15px 18px; text-align: center; color: #dc3545; font-weight: 700;">{{ item.absent }}</td>
                                    <td style="padding: 15px 18px; text-align: center;">
                                        <strong style="font-size: 1.1em;">{{ item.percent }}%</strong>
                                        <div style="height: 8px; background: #f0f0f0; border-radius: 4px; margin-top: 6px; overflow: hidden;">
                                            <div style="height: 100%; background: {% if item.percent >= 75 %}linear-gradient(90deg, #28a745, #20c997){% elif item.percent >= 65 %}linear-gradient(90deg, #ffc107, #ff9800){% else %}linear-gradient(90deg, #dc3545, #c82333){% endif %}; 
                                                        width: {{ item.percent }}%; transition: width 0.5s ease;"></div>
                                        </div>
                                    </td>
                                    <td style="padding: 15px 18px; text-align: center;">
                                        <span class="status-badge status-{{ item['status']|lower }}">
                                            {{ item.badge }} {{ item.status }}
                                        </span>
                                    </td>
                                </tr>
                            {% endfor %}
                        </tbody>
                    </table>
                </div>

                <h2 style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); -webkit-background-clip: text; -webkit-text-fill-color: transparent; background-clip: text; margin: 50px 0 30px 0; font-size: 1.8em; font-weight: 800;">📈 Analytics & Charts</h2>

                <div class="charts-grid">
                    <div class="chart-container">
                        <div class="chart-title">📊 Attendance % Line Graph</div>
                        <canvas id="percentChart" style="max-height: 300px;"></canvas>
                    </div>

                    <div class="chart-container">
                        <div class="chart-title">🎯 Status Distribution</div>
                        <canvas id="statusChart" style="max-height: 300px;"></canvas>
                    </div>

                    <div class="chart-container">
                        <div class="chart-title">✅ Present vs Absent</div>
                        <canvas id="presentAbsentChart" style="max-height: 300px;"></canvas>
                    </div>
                </div>
            </div>
        </div>

        <script>
            const students = {{ report|tojson }};
            const labels = students.map(s => s.roll);
            const percentages = students.map(s => s.percent);
            const present = students.map(s => s.present);
            const absent = students.map(s => s.absent);

            // Line Chart
            const ctxPercent = document.getElementById('percentChart').getContext('2d');
            new Chart(ctxPercent, {
                type: 'line',
                data: {
                    labels: labels,
                    datasets: [{
                        label: 'Attendance %',
                        data: percentages,
                        borderColor: '#667eea',
                        backgroundColor: 'rgba(102, 126, 234, 0.1)',
                        borderWidth: 4,
                        fill: true,
                        tension: 0.4,
                        pointBackgroundColor: '#667eea',
                        pointBorderColor: '#fff',
                        pointBorderWidth: 3,
                        pointRadius: 6,
                        pointHoverRadius: 8,
                        segment: {
                            borderDash: (ctx) => ctx.p0DataIndex % 2 === 0 ? [0] : [0]
                        }
                    }]
                },
                options: {
                    responsive: true,
                    maintainAspectRatio: true,
                    plugins: {
                        legend: { display: true, labels: { font: { size: 12, weight: 'bold' } } }
                    },
                    scales: {
                        y: {
                            beginAtZero: true,
                            max: 100,
                            ticks: { callback: function(value) { return value + '%'; } },
                            grid: { color: 'rgba(0, 0, 0, 0.05)' }
                        },
                        x: { grid: { display: false } }
                    }
                }
            });

            // Doughnut Chart
            const statusCounts = {
                Safe: students.filter(s => s.status === 'Safe').length,
                Warning: students.filter(s => s.status === 'Warning').length,
                Defaulter: students.filter(s => s.status === 'Defaulter').length
            };

            const ctxStatus = document.getElementById('statusChart').getContext('2d');
            new Chart(ctxStatus, {
                type: 'doughnut',
                data: {
                    labels: ['🟢 Safe (≥75%)', '🟡 Warning (65-74%)', '🔴 Defaulter (<65%)'],
                    datasets: [{
                        data: [statusCounts.Safe, statusCounts.Warning, statusCounts.Defaulter],
                        backgroundColor: ['#28a745', '#ffc107', '#dc3545'],
                        borderColor: '#fff',
                        borderWidth: 3
                    }]
                },
                options: {
                    responsive: true,
                    plugins: {
                        legend: { position: 'bottom', labels: { font: { size: 12, weight: 'bold' }, padding: 20 } }
                    }
                }
            });

            // Bar Chart
            const ctxBar = document.getElementById('presentAbsentChart').getContext('2d');
            new Chart(ctxBar, {
                type: 'bar',
                data: {
                    labels: labels,
                    datasets: [
                        {
                            label: '✅ Present',
                            data: present,
                            backgroundColor: '#28a745',
                            borderColor: '#228636',
                            borderWidth: 2
                        },
                        {
                            label: '❌ Absent',
                            data: absent,
                            backgroundColor: '#dc3545',
                            borderColor: '#bd2130',
                            borderWidth: 2
                        }
                    ]
                },
                options: {
                    responsive: true,
                    scales: {
                        x: { stacked: false },
                        y: { stacked: false }
                    },
                    plugins: {
                        legend: { labels: { font: { size: 12, weight: 'bold' } } }
                    }
                }
            });
        </script>
    </body>
    </html>
    """

    return render_template_string(html, 
                                 class_id=class_id,
                                 report=report,
                                 safe_count=safe_count,
                                 warning_count=warning_count,
                                 defaulter_count=defaulter_count)


@app.route("/report/student/<int:sid>")
def student_report(sid):
    if not session.get("admin"):
        return "Admin access required"

    db = get_db()

    student = db.execute(
        "SELECT roll_no, name FROM students WHERE id=?",
        (sid,)
    ).fetchone()

    records = db.execute(
        "SELECT date, status, od_reason FROM attendance WHERE student_id=? ORDER BY date",
        (sid,)
    ).fetchall()

    stats = get_student_attendance_stats(sid, db)

    absent_dates = [r[0] for r in records if r[1] == "Absent"]
    od_records = [(r[0], r[2]) for r in records if r[1] == "OD"]

    db.close()

    html = """
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <title>Student Report</title>
        """ + BASE_STYLES + """
    </head>
    <body>
        <div class="container">
            <div class="header">
                <h1 style="color: #667eea;">📊 Student Report</h1>
                <a href="/" class="btn btn-secondary">← Back</a>
            </div>

            <div style="background: rgba(255, 255, 255, 0.98); border-radius: 20px; padding: 40px; box-shadow: 0 20px 60px rgba(0, 0, 0, 0.1); border: 1px solid rgba(255, 255, 255, 0.5);">
                <div style="margin-bottom: 35px;">
                    <h2 style="color: #667eea; margin-bottom: 12px; font-size: 2em; font-weight: 800;">{{ student[1] }}</h2>
                    <p style="color: #666; font-size: 1.2em; font-weight: 600;">📋 Roll No: <strong style="color: #667eea;">{{ student[0] }}</strong></p>
                </div>

                <div class="stats-summary">
                    <div class="stat-box">
                        <h3>📊 Attendance %</h3>
                        <div class="value">{{ percent }}%</div>
                    </div>
                    <div class="stat-box" style="background: linear-gradient(135deg, {{ color }} 0%, {{ color }}cc 100%);">
                        <h3>{{ badge }} Status</h3>
                        <div class="value" style="font-size: 1.8em;">{{ status }}</div>
                    </div>
                </div>

                {% if absent_dates %}
                    <h3 style="color: #dc3545; margin: 35px 0 20px 0; font-size: 1.5em; font-weight: 800;">❌ Absent Dates ({{ absent_dates|length }})</h3>
                    <div style="background: linear-gradient(135deg, #fff5f5 0%, #ffe0e0 100%); padding: 25px; border-radius: 12px; margin-bottom: 30px; border-left: 5px solid #dc3545;">
                        <ul style="list-style: none; padding: 0;">
                            {% for d in absent_dates %}
                                <li style="padding: 12px 0; color: #721c24; font-weight: 600; border-bottom: 1px solid rgba(0,0,0,0.1);">📅 {{ d }}</li>
                            {% endfor %}
                        </ul>
                    </div>
                {% endif %}

                {% if od_records %}
                    <h3 style="color: #ffc107; margin: 35px 0 20px 0; font-size: 1.5em; font-weight: 800;">📝 OD Records ({{ od_records|length }})</h3>
                    <table style="width: 100%; border-collapse: collapse;">
                        <thead>
                            <tr style="background: linear-gradient(135deg, #ffc107 0%, #ff9800 100%); color: white;">
                                <th style="padding: 15px; text-align: left; font-weight: 700;">Date</th>
                                <th style="padding: 15px; text-align: left; font-weight: 700;">Reason</th>
                            </tr>
                        </thead>
                        <tbody>
                            {% for od_date, reason in od_records %}
                                <tr style="border-bottom: 2px solid #f0f0f0;">
                                    <td style="padding: 12px 15px; font-weight: 600;">{{ od_date }}</td>
                                    <td style="padding: 12px 15px; color: #666;">{{ reason or '-' }}</td>
                                </tr>
                            {% endfor %}
                        </tbody>
                    </table>
                {% endif %}
            </div>
        </div>
    </body>
    </html>
    """

    return render_template_string(html,
                                 student=student,
                                 percent=stats['percent'],
                                 status=stats['status'],
                                 badge=stats['badge'],
                                 color=stats['color'],
                                 absent_dates=absent_dates,
                                 od_records=od_records)


@app.route("/import/<int:class_id>", methods=["POST"])
def import_excel(class_id):
    if not session.get("admin"):
        flash("❌ Admin access required for importing students", "warning")
        return redirect(f"/attendance/{class_id}")

    file = request.files.get("file")
    if not file:
        flash("❌ No file selected", "warning")
        return redirect(f"/attendance/{class_id}")

    filename = file.filename.lower()
    students_data = []

    try:
        # Excel files
        if filename.endswith(('.xlsx', '.xls')):
            df = pd.read_excel(file)
            roll_col, name_col = detect_columns(df.columns)
            
            if not roll_col or not name_col:
                flash("❌ Required columns (Roll/Reg No & Name) not found in Excel", "warning")
                return redirect(f"/attendance/{class_id}")
            
            for _, row in df.iterrows():
                roll = row[roll_col]
                name = row[name_col]
                
                if pd.isna(roll) or pd.isna(name):
                    continue
                
                roll = str(roll).strip()
                name = str(name).strip()
                
                if not roll or not name or roll.lower() == "nan" or name.lower() == "nan":
                    continue
                
                students_data.append((roll, name))

        # Word files (DOCX)
        elif filename.endswith('.docx'):
            doc = docx.Document(file)
            
            for table in doc.tables:
                for row in table.rows[1:]:  # Skip header
                    if len(row.cells) >= 2:
                        roll = row.cells[0].text.strip()
                        name = row.cells[1].text.strip()
                        
                        if roll and name and roll.lower() != "nan" and name.lower() != "nan":
                            students_data.append((roll, name))
            
            # Also check paragraphs for data
            if not students_data:
                for para in doc.paragraphs:
                    text = para.text.strip()
                    if ' - ' in text or ',' in text:
                        parts = text.replace(' - ', ',').split(',')
                        if len(parts) >= 2:
                            roll = parts[0].strip()
                            name = parts[1].strip()
                            if roll and name:
                                students_data.append((roll, name))
            
            if not students_data:
                flash("❌ No student data found in Word document (use format: Roll - Name or Roll, Name)", "warning")
                return redirect(f"/attendance/{class_id}")

        # PDF files
        elif filename.endswith('.pdf'):
            pdf_reader = PdfReader(file)
            
            for page in pdf_reader.pages:
                text = page.extract_text()
                lines = text.split('\n')
                
                for line in lines:
                    line = line.strip()
                    if not line or line.lower() in ['roll no', 'name', 'roll', 'student']:
                        continue
                    
                    # Try to parse: Roll - Name or Roll, Name
                    if ' - ' in line:
                        parts = line.split(' - ')
                    elif ',' in line:
                        parts = line.split(',')
                    else:
                        continue
                    
                    if len(parts) >= 2:
                        roll = parts[0].strip()
                        name = parts[1].strip()
                        
                        if roll and name and len(roll) > 0 and len(name) > 0:
                            # Check if it looks like valid data (avoid headers/footers)
                            if not any(keyword in roll.lower() for keyword in ['roll', 'name', 'date', 'page']):
                                students_data.append((roll, name))
            
            if not students_data:
                flash("❌ No student data found in PDF (use format: Roll - Name per line)", "warning")
                return redirect(f"/attendance/{class_id}")

        else:
            flash("❌ Unsupported file format. Use Excel (.xlsx), Word (.docx), or PDF (.pdf)", "warning")
            return redirect(f"/attendance/{class_id}")

        # Remove duplicates while preserving order
        seen = set()
        unique_students = []
        for roll, name in students_data:
            key = (roll.lower(), name.lower())
            if key not in seen:
                seen.add(key)
                unique_students.append((roll, name))

        # Insert into database
        db = get_db()
        count = 0
        duplicates = 0

        for roll, name in unique_students:
            # Check if student already exists
            existing = db.execute(
                "SELECT id FROM students WHERE roll_no=? AND class_id=?",
                (roll, class_id)
            ).fetchone()
            
            if existing:
                duplicates += 1
                continue
            
            db.execute(
                "INSERT INTO students (roll_no, name, class_id) VALUES (?, ?, ?)",
                (roll, name, class_id)
            )
            count += 1

        db.commit()
        db.close()

        msg = f"✅ Import successful! {count} students added"
        if duplicates > 0:
            msg += f" ({duplicates} duplicates skipped)"
        
        flash(msg, "success")

    except Exception as e:
        flash(f"❌ Error importing file: {str(e)}", "warning")
        return redirect(f"/attendance/{class_id}")

    return redirect(f"/attendance/{class_id}")


@app.route("/export/total/excel/<int:class_id>")
def export_total_excel(class_id):
    if not session.get("admin"):
        return "Admin access required"

    db = get_db()

    students = db.execute(
        "SELECT id, roll_no, name FROM students WHERE class_id=? ORDER BY roll_no",
        (class_id,)
    ).fetchall()

    wb = Workbook()
    ws = wb.active
    ws.title = "Total Class Report"

    ws.append([
        "Roll No",
        "Name",
        "Total Days",
        "Present + OD",
        "Absent",
        "Attendance %",
        "Status"
    ])

    for sid, roll, name in students:
        stats = get_student_attendance_stats(sid, db)
        ws.append([
            roll,
            name,
            stats["total"],
            stats["present"],
            stats["absent"],
            stats["percent"],
            stats["status"]
        ])

    db.close()

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx")
    wb.save(tmp.name)

    return send_file(
        tmp.name,
        as_attachment=True,
        download_name="Total_Class_Attendance_Report.xlsx"
    )


@app.route("/export/day/excel/<int:class_id>")
def export_day_excel(class_id):
    if not session.get("admin"):
        return "Admin access required"

    report_date = request.args.get("date", date.today().isoformat())
    db = get_db()

    rows = db.execute("""
        SELECT s.roll_no, s.name, a.status, a.od_reason
        FROM attendance a
        JOIN students s ON a.student_id = s.id
        WHERE s.class_id=? AND a.date=?
    """, (class_id, report_date)).fetchall()

    db.close()

    wb = Workbook()
    ws = wb.active
    ws.title = "Day Report"

    ws.append(["Roll No", "Name", "Status", "OD Reason"])
    for r in rows:
        ws.append(r)

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx")
    wb.save(tmp.name)

    return send_file(
        tmp.name,
        as_attachment=True,
        download_name=f"Day_Report_{report_date}.xlsx"
    )


@app.route("/save_attendance", methods=["POST"])
def save_attendance():
    class_id = request.form["class_id"]
    selected_date = request.form["selected_date"]
    today = date.today().isoformat()
    db = get_db()

    if selected_date != today and not session.get("admin"):
        db.close()
        return "Admin password required to edit previous dates"

    db.execute("""
        DELETE FROM attendance
        WHERE date=? AND student_id IN
        (SELECT id FROM students WHERE class_id=?)
    """, (selected_date, class_id))

    students = db.execute(
        "SELECT id FROM students WHERE class_id=?", (class_id,)
    ).fetchall()

    defaulted = 0

    for (sid,) in students:
        status = request.form.get(f"status_{sid}")

        if status not in ["Present", "Absent", "OD"]:
            status = "Present"
            defaulted += 1

        reason = request.form.get(f"r_{sid}")

        db.execute(
            "INSERT INTO attendance (student_id, date, status, od_reason) VALUES (?, ?, ?, ?)",
            (sid, selected_date, status, reason)
        )

    db.commit()
    db.close()

    if defaulted > 0:
        flash(
            f"⚠️ {defaulted} students had no status selected. Marked as PRESENT.",
            "warning"
        )
    else:
        flash("✅ Attendance saved successfully!", "success")

    return redirect(f"/attendance/{class_id}?date={selected_date}")


def open_browser():
    webbrowser.open("http://127.0.0.1:5050")


if __name__ == "__main__":
    init_db_if_needed()

    if not os.environ.get("WERKZEUG_RUN_MAIN"):
        threading.Timer(1.2, open_browser).start()

    app.run(
        host="127.0.0.1",
        port=5050,
        debug=False,
        use_reloader=False
    )
